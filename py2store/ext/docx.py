from py2store.util import ModuleNotFoundErrorNiceMessage
from py2store.stores.local_store import LocalBinaryStore
from py2store import wrap_kvs

from io import BytesIO

with ModuleNotFoundErrorNiceMessage("docx wasn't found. Search and install python-docx package. "
                                    "For example, you could do: pip install python-docx"):
    import docx  # https://automatetheboringstuff.com/chapter13/


def get_text_from_docx(doc):
    """Get text from docx.Document object.
    More precisely, 'text' will be the newline-separated concatenation of the .text attributes of every paragraph.
    You can got a document object from a file path of pointer f by doing:
        import docx  # pip install python-docx
        doc = docx.Document(f)
    """
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def bytes_to_doc(doc_bytes):
    return docx.Document(BytesIO(doc_bytes))


LocalDocxStore = wrap_kvs(LocalBinaryStore, 'LocalDocxStore', obj_of_data=bytes_to_doc)

LocalDocxTextStore = wrap_kvs(LocalDocxStore, 'LocalDocxTextStore', obj_of_data=get_text_from_docx)
